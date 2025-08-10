"""
Resume Parser Module
Extracts text and key information from PDF and DOCX resumes
Enhanced with LangExtract for more accurate structured extraction
"""

import re
import os
import json
from typing import Dict, List, Optional
import PyPDF2
from docx import Document
import io
import google.generativeai as genai
import logging
import sys

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

try:
    from langextract import extract
    from langextract.data import ExampleData, Extraction, FormatType
    LANGEXTRACT_AVAILABLE = True
    logger.info("✅ LangExtract available for resume parsing")
except ImportError:
    LANGEXTRACT_AVAILABLE = False
    logger.warning("⚠️ LangExtract not available, using standard parsing methods")


class ResumeParser:
    """Parse resumes and extract key information for profile matching"""

    def __init__(self, use_gemini=True, use_langextract=True):
        self.use_gemini = use_gemini
        self.use_langextract = use_langextract and LANGEXTRACT_AVAILABLE

        # Initialize Gemini if API key is available
        if use_gemini:
            gemini_api_key = os.getenv("GEMINI_API_KEY")
            if gemini_api_key:
                genai.configure(api_key=gemini_api_key)
                self.model = genai.GenerativeModel("gemini-1.5-flash")
                self.gemini_api_key = gemini_api_key
                
                # Check if LangExtract is available
                if self.use_langextract:
                    logger.info("✨ LangExtract enabled for enhanced resume parsing")
            else:
                self.use_gemini = False
                self.use_langextract = False
                logger.warning("⚠️ Gemini API key not found, falling back to regex parsing")

        # Common section headers to identify
        self.section_headers = [
            "experience",
            "work experience",
            "professional experience",
            "employment",
            "education",
            "academic",
            "qualifications",
            "skills",
            "technical skills",
            "competencies",
            "projects",
            "achievements",
            "accomplishments",
            "certifications",
            "licenses",
            "military service",
            "service",
        ]

        # Keywords that indicate important connections
        self.connection_indicators = {
            "education": [
                "university",
                "college",
                "school",
                "academy",
                "institute",
                "degree",
                "bs",
                "ms",
                "phd",
                "mba",
            ],
            "military": [
                "military",
                "army",
                "navy",
                "air force",
                "marine",
                "coast guard",
                "veteran",
                "officer",
                "enlisted",
            ],
            "companies": [],  # Will be extracted dynamically
            "skills": [],  # Will be extracted dynamically
        }

    def parse_resume(self, file_content: bytes, filename: str) -> Dict:
        """
        Parse resume from file content
        Returns structured data about the resume
        """
        # Extract text based on file type
        text = ""
        if filename.lower().endswith(".pdf"):
            text = self._extract_pdf_text(file_content)
        elif filename.lower().endswith((".docx", ".doc")):
            text = self._extract_docx_text(file_content)
        else:
            text = file_content.decode("utf-8", errors="ignore")

        # Try LangExtract first if available
        if self.use_langextract and hasattr(self, "gemini_api_key"):
            try:
                logger.info(f"📄 Parsing resume '{filename}' with LangExtract...")
                parsed_data = self._parse_with_langextract(text)
                logger.info(f"✅ LangExtract parsing successful - Extracted {len(parsed_data.get('search_elements', []))} search elements")
                return parsed_data
            except Exception as e:
                logger.warning(f"⚠️ LangExtract parsing failed: {e}, falling back to Gemini")
                
        # Use Gemini for parsing if available, otherwise fall back to regex
        if self.use_gemini and hasattr(self, "model"):
            try:
                logger.info(f"🤖 Parsing resume with Gemini Flash...")
                parsed_data = self._parse_with_gemini(text)
                logger.info(f"✅ Gemini parsing successful - Extracted {len(parsed_data.get('search_elements', []))} search elements")
            except Exception as e:
                logger.warning(f"⚠️ Gemini parsing failed: {e}, falling back to regex")
                parsed_data = self._parse_text(text)
        else:
            logger.info("📝 Using regex-based parsing...")
            parsed_data = self._parse_text(text)

        return parsed_data

    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            return ""

    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            return text
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return ""

    def _parse_with_langextract(self, text: str) -> Dict:
        """Parse resume text using LangExtract for precise structured extraction"""
        
        logger.debug("🔍 Starting LangExtract extraction...")
        
        # Define extraction prompt
        prompt = """Extract the following structured information from this resume:
        
        1. email - Email address
        2. phone - Phone number  
        3. education - List of educational institutions with degree, field, and graduation year
        4. companies - List of companies worked at
        5. skills - Technical skills, tools, and programming languages
        6. military_service - Military branch, academy, rank if applicable
        7. certifications - Professional certifications
        8. locations - Cities/states lived or worked in
        9. achievements - Notable achievements or awards
        10. keywords - Industry keywords and specializations
        
        Return as structured JSON with these exact field names."""
        
        # Create example for better extraction
        example_text = """Jane Smith\njane@example.com\n(555) 987-6543\n\nEDUCATION\nHarvard University - BA Economics, 2015\n\nEXPERIENCE\nApple - Product Manager (2015-2020)\n\nSKILLS\nPython, SQL, Data Analysis"""
        
        example_extractions = [
            Extraction(extraction_class="email", extraction_text="jane@example.com"),
            Extraction(extraction_class="phone", extraction_text="(555) 987-6543"),
            Extraction(extraction_class="education", extraction_text="Harvard University - BA Economics, 2015"),
            Extraction(extraction_class="company", extraction_text="Apple"),
            Extraction(extraction_class="skill", extraction_text="Python"),
            Extraction(extraction_class="skill", extraction_text="SQL"),
            Extraction(extraction_class="skill", extraction_text="Data Analysis")
        ]
        
        example = ExampleData(
            text=example_text,
            extractions=example_extractions
        )
        
        # Perform extraction
        result = extract(
            text_or_documents=text,
            prompt_description=prompt,
            examples=[example],
            model_id="gemini-1.5-flash",
            api_key=self.gemini_api_key,
            format_type=FormatType.JSON,
            debug=False
        )
        
        # Parse the result - LangExtract returns AnnotatedDocument
        if hasattr(result, 'data'):
            extracted_data = result.data
        elif hasattr(result, 'extractions'):
            # Build data from extractions
            extracted_data = self._build_data_from_extractions(result.extractions)
        else:
            # Try to parse as JSON if it's a string result
            extracted_data = json.loads(str(result)) if isinstance(result, str) else {}
        
        # Calculate years of experience from work history
        years_experience = 0
        if "work_experience" in extracted_data:
            for exp in extracted_data["work_experience"]:
                if "duration" in exp:
                    # Try to extract years from duration string
                    duration = exp["duration"]
                    year_matches = re.findall(r"(\d+)\s*year", duration.lower())
                    if year_matches:
                        years_experience += int(year_matches[0])
        
        # Build the final parsed data structure
        parsed_data = {
            "full_text": text[:5000],
            "email": extracted_data.get("email"),
            "phone": extracted_data.get("phone"),
            "education": extracted_data.get("education", []),
            "companies": extracted_data.get("companies", []),
            "skills": extracted_data.get("skills", []),
            "military_service": extracted_data.get("military_service"),
            "years_experience": years_experience,
            "keywords": extracted_data.get("keywords", []),
            "certifications": extracted_data.get("certifications", []),
            "locations": extracted_data.get("locations", []),
            "achievements": extracted_data.get("achievements", []),
            "search_elements": []
        }
        
        # Build search elements with source grounding confidence
        search_elements = []
        
        # Add education elements with high confidence (source-grounded)
        for edu in parsed_data.get("education", []):
            if "institution" in edu and edu["institution"]:
                search_elements.append({
                    "category": "education",
                    "value": edu["institution"].lower(),
                    "weight": 35,  # Higher weight for source-grounded matches
                    "display": f"Alumni: {edu['institution']}",
                    "confidence": 0.95  # High confidence from LangExtract
                })
        
        # Add company elements
        for company in parsed_data.get("companies", []):
            if company:
                search_elements.append({
                    "category": "company", 
                    "value": company.lower(),
                    "weight": 30,
                    "display": f"Former {company}",
                    "confidence": 0.95
                })
        
        # Add military elements if present
        if parsed_data.get("military_service"):
            military = parsed_data["military_service"]
            if military.get("academy"):
                search_elements.append({
                    "category": "military",
                    "value": military["academy"].lower(),
                    "weight": 45,  # Very high weight for military academy
                    "display": f"{military['academy']} Alumni",
                    "confidence": 0.95
                })
            elif military.get("branch"):
                search_elements.append({
                    "category": "military",
                    "value": military["branch"].lower(),
                    "weight": 35,
                    "display": f"{military['branch']} Veteran",
                    "confidence": 0.95
                })
        
        # Add location elements
        for location in parsed_data.get("locations", [])[:3]:  # Top 3 locations
            if location:
                search_elements.append({
                    "category": "location",
                    "value": location.lower(),
                    "weight": 25,
                    "display": f"Connected to {location}",
                    "confidence": 0.9
                })
        
        # Add top skills (limit to avoid over-matching)
        for skill in parsed_data.get("skills", [])[:5]:  # Top 5 skills
            if skill:
                search_elements.append({
                    "category": "skill",
                    "value": skill.lower(),
                    "weight": 15,
                    "display": f"Shared skill: {skill}",
                    "confidence": 0.85
                })
        
        # Add certifications
        for cert in parsed_data.get("certifications", []):
            if cert:
                search_elements.append({
                    "category": "certification",
                    "value": cert.lower(),
                    "weight": 20,
                    "display": f"Both have {cert}",
                    "confidence": 0.9
                })
        
        # Add notable achievements
        for achievement in parsed_data.get("achievements", [])[:3]:
            if achievement and len(achievement) < 50:  # Short achievements as keywords
                search_elements.append({
                    "category": "achievement",
                    "value": achievement.lower(),
                    "weight": 20,
                    "display": achievement,
                    "confidence": 0.85
                })
        
        # Sort by weight and limit to top 15 elements
        search_elements.sort(key=lambda x: x["weight"], reverse=True)
        parsed_data["search_elements"] = search_elements[:15]
        
        # Add extraction metadata
        parsed_data["extraction_method"] = "langextract"
        parsed_data["extraction_confidence"] = "high"
        
        return parsed_data
    
    def _build_data_from_extractions(self, extractions: List) -> Dict:
        """Build structured data from LangExtract extractions"""
        data = {
            "email": None,
            "phone": None,
            "education": [],
            "companies": [],
            "skills": [],
            "military_service": None,
            "certifications": [],
            "locations": [],
            "achievements": [],
            "keywords": []
        }
        
        for ext in extractions:
            if ext.extraction_class == "email":
                data["email"] = ext.extraction_text
            elif ext.extraction_class == "phone":
                data["phone"] = ext.extraction_text
            elif ext.extraction_class == "education":
                # Parse education text
                edu_text = ext.extraction_text
                parts = edu_text.split("-")
                if len(parts) >= 2:
                    data["education"].append({
                        "institution": parts[0].strip(),
                        "degree": parts[1].strip() if len(parts) > 1 else ""
                    })
            elif ext.extraction_class == "company":
                data["companies"].append(ext.extraction_text)
            elif ext.extraction_class == "skill":
                data["skills"].append(ext.extraction_text)
            elif ext.extraction_class == "certification":
                data["certifications"].append(ext.extraction_text)
            elif ext.extraction_class == "location":
                data["locations"].append(ext.extraction_text)
            elif ext.extraction_class == "achievement":
                data["achievements"].append(ext.extraction_text)
            elif ext.extraction_class == "keyword":
                data["keywords"].append(ext.extraction_text)
            elif ext.extraction_class == "military":
                if not data["military_service"]:
                    data["military_service"] = {}
                # Parse military text for branch/academy/rank
                mil_text = ext.extraction_text.lower()
                if "air force" in mil_text or "usafa" in mil_text:
                    data["military_service"]["branch"] = "air force"
                    if "usafa" in mil_text or "academy" in mil_text:
                        data["military_service"]["academy"] = "USAFA"
                elif "army" in mil_text or "west point" in mil_text:
                    data["military_service"]["branch"] = "army"
                    if "west point" in mil_text:
                        data["military_service"]["academy"] = "West Point"
                elif "navy" in mil_text or "naval academy" in mil_text:
                    data["military_service"]["branch"] = "navy"
                    if "naval academy" in mil_text:
                        data["military_service"]["academy"] = "Naval Academy"
                elif "marine" in mil_text:
                    data["military_service"]["branch"] = "marines"
                elif "coast guard" in mil_text:
                    data["military_service"]["branch"] = "coast guard"
                    
                data["military_service"]["veteran"] = True
        
        return data

    def _parse_with_gemini(self, text: str) -> Dict:
        """Parse resume text using Gemini Flash LLM for better extraction"""

        prompt = (
            """Analyze this resume and extract the following information in JSON format:

{
  "email": "email address if found",
  "phone": "phone number if found",
  "education": [
    {
      "institution": "school/university name",
      "degree": "degree type (BS, MS, PhD, etc.)",
      "field": "field of study",
      "graduation_year": "year if found"
    }
  ],
  "companies": ["list of company names worked at"],
  "skills": ["list of technical skills, tools, programming languages"],
  "military_service": {
    "branch": "branch of service if any",
    "academy": "military academy if attended",
    "rank": "rank if mentioned",
    "veteran": true/false
  },
  "certifications": ["list of professional certifications"],
  "keywords": ["important industry keywords, roles, specializations"],
  "locations": ["cities/states lived or worked in"],
  "years_experience": estimated total years of experience (number),
  "search_elements": [
    {
      "category": "education/company/military/location/skill/achievement/keyword",
      "value": "the actual value to search for",
      "weight": importance score 1-50,
      "display": "display text for UI"
    }
  ]
}

Create search_elements for the MOST IMPORTANT networking connections. Prioritize:
1. Universities/Schools (weight: 35-40) - Alumni connections are strongest
2. Companies worked at (weight: 30-35) - Former colleagues matter
3. Military academies/service (weight: 40-45) - Very strong bonds
4. Hometowns/Cities lived in (weight: 25-30) - Geographic connections
5. Only 2-3 HIGH-LEVEL skills/interests (weight: 15-20) - e.g., "AI", "Machine Learning", not specific tools
6. Major life experiences (weight: 20-25) - e.g., "Founder", "CEO", significant achievements

IMPORTANT: Limit to 10-15 total search elements. Focus on BIG connections that create instant rapport, not minor technical details. Avoid specific programming languages or tools unless they define someone's career (e.g., "Python expert" not just "Python")

Resume text:
"""
            + text[:10000]
        )  # Limit to 10k chars for API

        try:
            response = self.model.generate_content(prompt)
            # Parse the JSON response
            json_str = response.text
            # Clean up the response if it has markdown code blocks
            if "```json" in json_str:
                json_str = json_str.split("```json")[1].split("```")[0]
            elif "```" in json_str:
                json_str = json_str.split("```")[1].split("```")[0]

            parsed_data = json.loads(json_str)

            # Ensure all required fields exist
            parsed_data.setdefault("full_text", text[:5000])
            parsed_data.setdefault("email", None)
            parsed_data.setdefault("phone", None)
            parsed_data.setdefault("education", [])
            parsed_data.setdefault("companies", [])
            parsed_data.setdefault("skills", [])
            parsed_data.setdefault("military_service", None)
            parsed_data.setdefault("years_experience", 0)
            parsed_data.setdefault("keywords", [])
            parsed_data.setdefault("certifications", [])
            parsed_data.setdefault("search_elements", [])

            # If no search elements were created, build them from extracted data
            if not parsed_data["search_elements"]:
                parsed_data["search_elements"] = self._build_search_elements(
                    parsed_data
                )

            return parsed_data

        except Exception as e:
            print(f"Error in Gemini parsing: {e}")
            raise

    def _parse_text(self, text: str) -> Dict:
        """Parse resume text and extract structured information"""
        # Clean and normalize text
        text = text.lower()

        # Extract key information
        data = {
            "full_text": text[:5000],  # Store first 5000 chars for reference
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "education": self._extract_education(text),
            "companies": self._extract_companies(text),
            "skills": self._extract_skills(text),
            "military_service": self._extract_military(text),
            "years_experience": self._estimate_experience(text),
            "keywords": self._extract_keywords(text),
            "certifications": self._extract_certifications(text),
            "search_elements": [],  # This will be used for matching
        }

        # Build search elements for matching
        data["search_elements"] = self._build_search_elements(data)

        return data

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text"""
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        match = re.search(email_pattern, text)
        return match.group(0) if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text"""
        phone_pattern = r"[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}"
        match = re.search(phone_pattern, text)
        return match.group(0) if match else None

    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education information"""
        education = []

        # Common university patterns
        university_patterns = [
            r"([a-z\s]+university)",
            r"([a-z\s]+college)",
            r"([a-z\s]+institute)",
            r"([a-z\s]+academy)",
            r"(usafa|usma|usna|uscga|usmma)",  # Military academies
        ]

        # Common degrees
        degree_patterns = [
            r"(bachelor|b\.s\.|bs\.|b\.a\.|ba\.)",
            r"(master|m\.s\.|ms\.|m\.a\.|ma\.|mba)",
            r"(ph\.d\.|phd|doctorate)",
        ]

        for pattern in university_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                education.append({"institution": match.strip(), "type": "university"})

        for pattern in degree_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                education.append({"degree": match.strip(), "type": "degree"})

        return education

    def _extract_companies(self, text: str) -> List[str]:
        """Extract company names from work experience"""
        companies = []

        # Look for patterns like "at Company" or "Company Name"
        # This is simplified - in production, you'd use NER
        company_patterns = [
            r"at\s+([A-Z][A-Za-z\s&]+)",
            r"[\n•]\s*([A-Z][A-Za-z\s&]+)\s*[\n•]",
            r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)\s*\|",
        ]

        for pattern in company_patterns:
            matches = re.findall(pattern, text[:2000])  # Focus on first part of resume
            companies.extend([m.strip() for m in matches if len(m.strip()) > 2])

        # Also look for known tech companies
        tech_companies = [
            "google",
            "meta",
            "facebook",
            "amazon",
            "apple",
            "microsoft",
            "openai",
            "anthropic",
            "tesla",
            "spacex",
            "uber",
            "airbnb",
            "stripe",
            "coinbase",
            "dropbox",
            "slack",
            "zoom",
            "oracle",
            "ibm",
            "intel",
            "nvidia",
            "amd",
            "salesforce",
            "adobe",
            "c3.ai",
            "c3 ai",
            "palantir",
            "databricks",
            "snowflake",
        ]

        for company in tech_companies:
            if company in text:
                companies.append(company)

        return list(set(companies))  # Remove duplicates

    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills"""
        skills = []

        # Programming languages
        languages = [
            "python",
            "java",
            "javascript",
            "typescript",
            "c++",
            "c#",
            "ruby",
            "go",
            "rust",
            "swift",
            "kotlin",
            "scala",
            "r",
        ]

        # Frameworks and tools
        frameworks = [
            "react",
            "angular",
            "vue",
            "django",
            "flask",
            "fastapi",
            "spring",
            "nodejs",
            "express",
            "rails",
            "laravel",
            "tensorflow",
            "pytorch",
            "scikit-learn",
            "pandas",
            "numpy",
        ]

        # Cloud and DevOps
        cloud_tools = [
            "aws",
            "azure",
            "gcp",
            "google cloud",
            "kubernetes",
            "docker",
            "terraform",
            "jenkins",
            "gitlab",
            "github",
            "ci/cd",
        ]

        # Databases
        databases = [
            "sql",
            "postgresql",
            "mysql",
            "mongodb",
            "redis",
            "elasticsearch",
            "cassandra",
            "dynamodb",
            "firebase",
        ]

        all_skills = languages + frameworks + cloud_tools + databases

        for skill in all_skills:
            if skill in text:
                skills.append(skill)

        return skills

    def _extract_military(self, text: str) -> Optional[Dict]:
        """Extract military service information"""
        # Check for negative context first
        negative_patterns = [
            "no military",
            "not military",
            "non-military",
            "without military",
            "never served",
            "civilian",
        ]

        for pattern in negative_patterns:
            if pattern in text:
                return None

        military_keywords = [
            "air force",
            "army",
            "navy",
            "marine",
            "coast guard",
            "military",
            "veteran",
            "active duty",
            "reserve",
            "national guard",
            "usafa",
            "west point",
            "usma",
            "naval academy",
            "usna",
            "officer",
            "enlisted",
            "sergeant",
            "lieutenant",
            "captain",
            "major",
            "colonel",
        ]

        military_info = {}
        for keyword in military_keywords:
            if keyword in text:
                if "branch" not in military_info:
                    if any(
                        branch in keyword
                        for branch in [
                            "air force",
                            "army",
                            "navy",
                            "marine",
                            "coast guard",
                        ]
                    ):
                        military_info["branch"] = keyword
                military_info["veteran"] = True

                # Check for academy
                if any(
                    academy in keyword
                    for academy in [
                        "usafa",
                        "usma",
                        "usna",
                        "west point",
                        "naval academy",
                    ]
                ):
                    military_info["academy"] = keyword

        return military_info if military_info else None

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract professional certifications"""
        certifications = []

        cert_patterns = [
            "aws certified",
            "azure certified",
            "google certified",
            "pmp",
            "cissp",
            "ccna",
            "ccnp",
            "comptia",
            "scrum master",
            "six sigma",
            "itil",
            "cpa",
            "cfa",
        ]

        for cert in cert_patterns:
            if cert in text:
                certifications.append(cert)

        return certifications

    def _estimate_experience(self, text: str) -> int:
        """Estimate years of experience from resume"""
        # Look for year patterns (e.g., 2018-2023)
        year_pattern = r"\b(19\d{2}|20\d{2})\b"
        years = re.findall(year_pattern, text)

        if years:
            years = [int(y) for y in years]
            # Rough estimate: difference between earliest and latest year
            return max(years) - min(years)

        # Look for explicit mentions
        exp_pattern = r"(\d+)\+?\s*years?\s*(?:of\s*)?experience"
        match = re.search(exp_pattern, text)
        if match:
            return int(match.group(1))

        return 0

    def _extract_keywords(self, text: str) -> List[str]:
        """Extract important keywords for matching"""
        # This would ideally use TF-IDF or similar
        # For now, extract nouns and important terms
        keywords = []

        # Industry keywords
        industries = [
            "fintech",
            "healthtech",
            "edtech",
            "biotech",
            "cleantech",
            "saas",
            "b2b",
            "b2c",
            "marketplace",
            "platform",
            "ai",
            "machine learning",
            "data science",
            "analytics",
            "blockchain",
            "crypto",
            "web3",
            "defi",
        ]

        # Role keywords
        roles = [
            "engineer",
            "developer",
            "architect",
            "manager",
            "director",
            "analyst",
            "scientist",
            "designer",
            "consultant",
            "lead",
            "senior",
            "principal",
            "staff",
            "founder",
            "ceo",
            "cto",
        ]

        all_keywords = industries + roles

        for keyword in all_keywords:
            if keyword in text:
                keywords.append(keyword)

        return keywords

    def _build_search_elements(self, data: Dict) -> List[Dict]:
        """
        Build search elements that will be used for matching profiles
        Each element has a category, value, and weight for scoring
        """
        elements = []

        # Add education elements
        for edu in data.get("education", []):
            if "institution" in edu:
                elements.append(
                    {
                        "category": "education",
                        "value": edu["institution"],
                        "weight": 30,  # High weight for education matches
                        "display": f"Alumni: {edu['institution'].title()}",
                    }
                )

        # Add company elements
        for company in data.get("companies", []):
            elements.append(
                {
                    "category": "company",
                    "value": company.lower(),
                    "weight": 25,  # Good weight for company matches
                    "display": f"Former {company.title()}",
                }
            )

        # Add military elements
        if data.get("military_service"):
            military = data["military_service"]
            if military.get("academy"):
                elements.append(
                    {
                        "category": "military",
                        "value": military["academy"],
                        "weight": 40,  # Very high weight for military academy
                        "display": f"{military['academy'].upper()} Alumni",
                    }
                )
            elif military.get("branch"):
                elements.append(
                    {
                        "category": "military",
                        "value": military["branch"],
                        "weight": 30,
                        "display": f"{military['branch'].title()} Veteran",
                    }
                )

        # Add skill elements (lower weight)
        for skill in data.get("skills", [])[:10]:  # Top 10 skills
            elements.append(
                {
                    "category": "skill",
                    "value": skill,
                    "weight": 10,
                    "display": f"Shared skill: {skill}",
                }
            )

        # Add certification elements
        for cert in data.get("certifications", []):
            elements.append(
                {
                    "category": "certification",
                    "value": cert,
                    "weight": 15,
                    "display": f"Both have {cert.upper()}",
                }
            )

        # Add keyword elements
        for keyword in data.get("keywords", [])[:5]:  # Top 5 keywords
            elements.append(
                {
                    "category": "keyword",
                    "value": keyword,
                    "weight": 5,
                    "display": keyword.title(),
                }
            )

        return elements


def parse_resume_file(
    file_content: bytes, filename: str, use_gemini: bool = True, use_langextract: bool = True
) -> Dict:
    """Convenience function to parse a resume file

    Args:
        file_content: The binary content of the file
        filename: The name of the file
        use_gemini: Whether to use Gemini Flash for parsing (default: True)
        use_langextract: Whether to use LangExtract for parsing (default: True)
    """
    parser = ResumeParser(use_gemini=use_gemini, use_langextract=use_langextract)
    return parser.parse_resume(file_content, filename)
